import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('data/Jeevan_KK_Data_Engineer_DE.docx')

def set_single(para, text, bold=False, size=None):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = para.add_run(text)
    r.bold = bold
    if size: r.font.size = size

def set_two(para, t1, b1, t2, b2):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r1 = para.add_run(t1); r1.bold = b1
    r2 = para.add_run(t2); r2.bold = b2

p = doc.paragraphs

# [1] Subtitle
set_single(p[1],
    'Senior Data Platform Engineer | Python · Airflow · Spark · Kafka · HDFS · Iceberg',
    bold=False, size=152400)

# [5] Summary line 1
set_single(p[5],
    'Senior Data Platform / DataOps Engineer with 8+ years of experience building, operating, and migrating large-scale data platforms. Deep hands-on expertise in Apache Airflow production ownership: 40+ DAGs, scheduler internals, executors, sensors, pools, SLA tracking, callbacks/listeners, and automated backfill/retry flows. Experienced with the open-source stack (Python, PySpark, Kafka, HDFS, Apache Iceberg, Trino/Presto, dbt) across batch and streaming workloads at tens of millions of events per day.')

# [6] Summary line 2
set_single(p[6],
    'Proven track record migrating legacy ETL environments to modern stacks, with a platform-reliability mindset: monitoring, alerting, runbooks, incident diagnostics, and DataOps automation as first-class deliverables. Experienced leading engineers through code reviews and engineering standards; comfortable in remote, distributed, and cross-functional team structures.')

# [9] Big Data — Airflow + open-source stack first
set_two(p[9], 'Big Data & Processing: ', True,
    'Apache Airflow (DAGs, KubernetesExecutor/CeleryExecutor, sensors, pools, SLAs, callbacks/listeners, backfills, retries) · Apache Spark, PySpark, Spark Structured Streaming · Apache Kafka · Apache Flink · HDFS · Apache Iceberg · Trino / Presto · Databricks · Amazon EMR · Hadoop',
    False)

# [10] Data Warehouses — Iceberg/open-source first
set_two(p[10], 'Data Warehouses & Databases: ', True,
    'Apache Iceberg · Delta Lake · Amazon Redshift · Snowflake · PostgreSQL · Oracle · SQL Server · MySQL · DynamoDB',
    False)

# [11] ETL — Airflow/dbt/migration first
set_two(p[11], 'ETL & Data Integration: ', True,
    'Apache Airflow · dbt · AWS Glue · Great Expectations · CDC patterns (incremental loads, high-watermark, DMS) · legacy-to-modern pipeline migration · AWS DMS · SSIS · AppFlow',
    False)

# [14] DevOps — add Kubernetes, Grafana/Prometheus
set_two(p[14], 'DevOps & Governance: ', True,
    'Kubernetes (deployments, jobs, cronjobs, configs, secrets, resource limits) · Docker · Git · GitHub Actions · GitLab CI · CI/CD · Terraform (IaC) · Grafana / Prometheus-style observability · GDPR-compliant pipelines · SOX · RBAC',
    False)

# [49] Kafka project title
set_two(p[49], 'IT Services / Streaming Analytics — ', True,
    'Real-Time Kafka Platform — 50M+ Events/Day, Observability & DataOps',
    True)

# [50] Kafka bullet 1 — lag/health monitoring
set_single(p[50],
    'Designed and implemented a Kafka-based real-time ingestion pipeline handling > 50M events per day at sub-second latency; set up consumer lag monitoring, topic health dashboards, and alerting for throughput anomalies and partition imbalance.')

# [52] Kafka bullet 3 — observability
set_single(p[52],
    'Persisted processed data into AWS S3, Redshift, and DynamoDB; implemented end-to-end observability via CloudWatch and Datadog covering pipeline failures, queue depth, data freshness, and completeness SLOs.')

# [53] Kafka bullet 4 — reliability/DataOps
set_single(p[53],
    'Engineered for data consistency, availability, and fault-tolerance under high-throughput workloads; authored runbooks for incident diagnostics and postmortems across Kafka, Spark, and downstream systems.')

# [66] Airflow project title
set_two(p[66], 'IT Services / Orchestration — ', True,
    'Production Airflow Platform — DAG Engineering, Observability & DataOps Automation',
    True)

# [67] Airflow bullet 1 — internals depth
set_single(p[67],
    'Designed and operated 40+ Airflow DAGs in production (data ingestion, transformation, validation) with > 99% task success rate. Configured executors (CeleryExecutor / KubernetesExecutor), worker pools, task queues, and priority weights for resource isolation and throughput optimisation.')

# [68] Airflow bullet 2 — sensors, SLAs, callbacks
set_single(p[68],
    'Implemented SLA tracking and callback/listener hooks (on_failure_callback, on_retry_callback, SLA miss alerts) alongside custom sensors for upstream data arrival, file readiness, and API availability, eliminating hard-coded waits and reducing pipeline failures by > 40%.')

# [69] Airflow bullet 3 — automation/CI/CD
set_single(p[69],
    'Automated operational workflows: templated DAG deployment scripts, backfill/retry/recovery CLIs, pre-release validation checks, and CI/CD via GitHub Actions, enabling zero-downtime DAG releases and eliminating manual deployment steps.')

# [70] Airflow bullet 4 — monitoring/runbooks
set_single(p[70],
    'Set up Airflow observability: scheduler heartbeat monitoring, task-duration anomaly alerting, queue-time dashboards, and dead-letter handling; authored runbooks and incident postmortems to reduce recurring failures and operational toil.')

# [83] GE project — reframe as DataOps/governance
set_two(p[83], 'Pharma / Life Sciences — ', True,
    'DataOps & Data Governance — Data Quality, Lineage & Observability Platform',
    True)

# [84] GE bullet 1 — governance/ownership
set_single(p[84],
    'Defined and enforced data quality contracts (missing values, duplicates, schema drift, referential integrity) across 100+ datasets using Great Expectations, publishing expectation suites as shared ownership artefacts with named data owners and metadata annotations.')

# [85] GE bullet 2 — SLO dashboards
set_single(p[85],
    'Built automated data profiling, freshness checks, and SLO/SLA dashboards providing data consumers with visibility into pipeline health and completeness without manual intervention.')

# [86] GE bullet 3 — alerting/MTTR
set_single(p[86],
    'Wired failed-check alerts into incident channels with structured diagnostics: failing expectations, affected datasets, upstream lineage, and recommended remediation steps, reducing mean time to resolution (MTTR) for data incidents.')

# [87] GE bullet 4 — integration/privacy-by-design
set_single(p[87],
    'Integrated the framework with Airflow (operator-level hooks), Spark, and Glue for end-to-end pipeline validation; implemented privacy-by-design checks for PII field presence and data access audit trails.')

# [60] Lakehouse project — HDFS/Iceberg/Trino focus
set_two(p[60], 'IT Services / Data Platform — ', True,
    'Open-Source Data Lakehouse — HDFS, Apache Iceberg, Trino & Delta Lake',
    True)

# [61] Lakehouse bullet 1 — HDFS/Iceberg zones
set_single(p[61],
    'Architected a multi-zone Data Lakehouse (Raw / Curated / Serving) on HDFS and AWS S3 using Apache Iceberg and Delta Lake for ACID-compliant analytical workloads, supporting time-travel, schema evolution, and partition evolution without table rewrites.')

# [63] Lakehouse bullet 3 — metadata/lineage/governance
set_single(p[63],
    'Implemented data cataloguing, metadata management, and lineage tracking, establishing naming conventions, ownership records, and access-pattern auditing to support data governance and privacy-by-design requirements.')

# [65] Lakehouse bullet 5 — Trino query optimisation
set_single(p[65],
    'Optimised analytical query performance using Trino / Presto and Athena with predicate pushdown, file compaction (Iceberg OPTIMIZE), and partition pruning, reducing ad-hoc query latency by > 50%.')

out = 'data/resumes/Jeevan-Kondasingu-MyGames-SeniorDataPlatformEngineer.docx'
doc.save(out)
print(f'Saved: {out}')
