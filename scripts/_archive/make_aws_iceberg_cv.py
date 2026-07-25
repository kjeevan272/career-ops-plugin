import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('data/Jeevan_KK_Data_Engineer_DE.docx')

def set_single(para, text, bold=False, size=None):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = para.add_run(text)
    r.bold = bold
    if size: r.font.size = size

def set_two(para, t1, b1, t2, b2):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r1 = para.add_run(t1); r1.bold = b1
    r2 = para.add_run(t2); r2.bold = b2

p = doc.paragraphs

# [1] Subtitle - AWS Iceberg architect focus
set_single(p[1],
    'AWS Iceberg Data Engineer and Architect | S3, EMR, Glue, Lambda, Redshift, PySpark, Apache Iceberg',
    bold=False, size=152400)

# [5] Summary line 1 - AWS native + Iceberg internals + PySpark scale
set_single(p[5],
    'Senior Data Engineer and Data Architect with 8+ years of experience designing and operating '
    'scalable, high-performance data platforms on AWS. Deep, hands-on expertise with Apache Iceberg '
    'lakehouses on Amazon S3, including hidden and evolving partitioning, schema evolution, snapshot '
    'based time travel, copy-on-write and merge-on-read tables, and compaction tuning. Strong command '
    'of AWS native services (S3, EMR, Glue, Lambda, Redshift, Athena, Step Functions, Kinesis, DMS, '
    'Lake Formation) and of PySpark and Spark for terabyte scale batch and streaming workloads '
    'handling hundreds of millions of rows per day.')

# [6] Summary line 2 - outcomes, modeling, CI/CD, leadership, certs, agentic
set_single(p[6],
    'Owned end-to-end architecture from ingestion through ETL and ELT to analytics ready Iceberg and '
    'Redshift data models, with measurable outcomes including a 40% reduction in pipeline processing '
    'time, around 30% lower compute and storage cost, and over 99% data quality scores. Experienced '
    'in data modeling, CI/CD and DevOps automation, and the performance, reliability, and security of '
    'cloud data platforms, alongside technical leadership and stakeholder management. Certified AWS '
    'Solutions Architect Associate, Databricks Developer for Apache Spark, and Snowflake SnowPro Core, '
    'with hands-on exposure to agentic and LLM driven data workflows.')

# [8] Cloud and Infrastructure - AWS heavy
set_two(p[8], 'Cloud and Infrastructure: ', True,
    'AWS (S3, EMR, Glue, Glue Data Catalog, Lambda, Redshift, Redshift Spectrum, Athena, Step Functions, '
    'Kinesis, DMS, Lake Formation, CloudWatch, IAM, KMS, CodePipeline), Azure (ADF, ADLS Gen2, '
    'Databricks), GCP (BigQuery)',
    False)

# [9] Big Data and Processing - Iceberg internals lead
set_two(p[9], 'Big Data and Processing: ', True,
    'Apache Iceberg (hidden partitioning, partition and schema evolution, time travel, copy-on-write '
    'and merge-on-read, compaction), Apache Spark, PySpark, Spark SQL, Spark Structured Streaming, '
    'Amazon EMR, Apache Kafka, Apache Flink, Apache Airflow (MWAA), Delta Lake, Hadoop, HDFS',
    False)

# [10] Data Warehouses and Databases
set_two(p[10], 'Data Warehouses and Databases: ', True,
    'Apache Iceberg, Amazon Redshift, Snowflake, Delta Lake, AWS Glue Data Catalog, PostgreSQL, Oracle, '
    'SQL Server, MySQL, DynamoDB, Google BigQuery',
    False)

# [11] ETL and Data Integration
set_two(p[11], 'ETL and Data Integration: ', True,
    'AWS Glue (PySpark), Amazon EMR Spark, Iceberg MERGE INTO upserts and CDC, AWS DMS, dbt, '
    'Great Expectations, AppFlow, Salesforce Bulk and REST and Streaming API, Lambda and EventBridge triggers',
    False)

# [14] DevOps and Governance - CI/CD emphasis
set_two(p[14], 'DevOps and Governance: ', True,
    'Git, GitHub Actions, AWS CodePipeline, GitLab CI, CI/CD, Terraform (IaC), Docker, GDPR compliant '
    'pipelines, SOX, RBAC, IAM and KMS security',
    False)

# [15] BI and Visualisation
set_two(p[15], 'BI and Visualisation: ', True,
    'Power BI, Amazon QuickSight, Tableau, Pandas, Matplotlib, Seaborn',
    False)

# [16] ML and AI - add agentic workflows (good to have)
set_two(p[16], 'ML and AI: ', True,
    'scikit-learn, TensorFlow (integrated in data pipelines), Predictive Analytics, KNN, SVM, Random '
    'Forest, plus exposure to agentic and LLM driven data workflows and automation',
    False)

# [19] Date / location line
set_single(p[19], '07/2017 to Present, New Barnet, United Kingdom')

# [20] Role / company line
set_single(p[20], 'Data Engineer, Global Research and Innovation Centre (UK) Limited', bold=True)

# [21] Role description - AWS Iceberg architecture leadership
set_single(p[21],
    'Designing, building, and operating scalable AWS native batch and streaming data platforms for '
    'international clients, with a primary focus on Apache Iceberg lakehouses on Amazon S3. Leading '
    'end-to-end delivery from architecture and data modeling through PySpark pipeline build, '
    'performance tuning, governance, and stakeholder delivery, and providing technical leadership and '
    'best practices for cloud data architecture handling hundreds of millions of rows.')

# [35] E-Commerce EDW header - reframe as DW + Lakehouse
set_two(p[35], 'E-Commerce / Online Retail: ', True,
    'Enterprise Data Warehouse and Lakehouse for E-Commerce BI', True)
set_single(p[36],
    'Led the build of serverless ETL pipelines with AWS Glue Jobs, the Glue Data Catalog, and PySpark '
    'for batch and incremental integration of heterogeneous sources, processing over 200M records per day.')
set_single(p[37],
    'Designed and tuned PySpark workloads (partitioning, schema evolution, complex aggregations, Spark '
    'SQL) on terabyte scale datasets, reducing cluster cost by around 25%.')

# [42] Insurance Databricks/EMR -> EMR + Iceberg real-time platform
set_two(p[42], 'Insurance Analytics: ', True,
    'Real-Time EMR and Iceberg Data Platform with Glue and Kinesis', True)
set_single(p[43],
    'Architected and operated an end-to-end real-time data pipeline on Amazon EMR, AWS Glue, and '
    'Kinesis for insurance telemetry and claims data, writing to Apache Iceberg tables on S3 for ACID '
    'guarantees and incremental reads.')
set_single(p[44],
    'Ingested raw events from Kinesis and S3 and developed custom PySpark ETL jobs producing Iceberg '
    'and Parquet outputs with idempotent, exactly once MERGE based loads.')
set_single(p[45],
    'Implemented Change Data Capture for incremental loads with guaranteed consistency between source '
    'and destination, using Iceberg snapshots for auditability and rollback.')
set_single(p[46],
    'Tuned Spark jobs on EMR clusters (data skew, partitioning, executor sizing, and Iceberg '
    'compaction), reducing processing time by around 35% and significantly cutting infrastructure cost.')
set_single(p[47],
    'Delivered CI/CD pipelines for EMR and Glue jobs and Databricks notebooks via GitHub Actions and '
    'AWS CodePipeline, with automated testing and deployment.')

# [60] FLAGSHIP - Apache Iceberg Lakehouse on AWS
set_two(p[60], 'Insurance and Retail Analytics: ', True,
    'Apache Iceberg Lakehouse on AWS S3, Glue, EMR, Athena, and Redshift', True)
set_single(p[61],
    'Architected and built a petabyte scale Apache Iceberg lakehouse on Amazon S3 with the AWS Glue '
    'Data Catalog as the Iceberg catalog, serving ACID compliant analytical workloads across EMR '
    'Spark, Athena, and Redshift Spectrum.')
set_single(p[62],
    'Modeled Iceberg tables for high write throughput using hidden partitioning and partition '
    'evolution, tuning target file sizes and running periodic compaction with rewrite_data_files and '
    'expire_snapshots to keep query latency low on tables with hundreds of millions of rows.')
set_single(p[63],
    'Implemented schema evolution and snapshot based time travel for safe backfills, reproducible '
    'reporting, and incremental MERGE INTO upserts, choosing copy-on-write or merge-on-read per table '
    'based on read and write patterns.')
set_single(p[64],
    'Built ingestion and transformation pipelines in PySpark on Amazon EMR and AWS Glue, landing Raw, '
    'Bronze, Silver, and Gold layers and enforcing RBAC and encryption through AWS Lake Formation, '
    'IAM, and KMS.')
set_single(p[65],
    'Optimised query performance through manifest and partition pruning, Athena and Redshift Spectrum '
    'external tables, and regular metadata maintenance, cutting analyst query runtimes significantly '
    'while lowering S3 storage cost with lifecycle policies.')

# [66] Orchestration - CI/CD and DevOps emphasis
set_two(p[66], 'IT Services / Orchestration: ', True,
    'ETL Orchestration and CI/CD Automation with Apache Airflow and MWAA', True)
set_single(p[67],
    'Developed 40+ Airflow DAGs for ingestion, transformation, and validation, achieving over 99% '
    'success rate in production, with retries, SLAs, and alerting.')
set_single(p[70],
    'Integrated Airflow with AWS Glue, EMR, Iceberg, Snowflake, and Redshift, and automated DAG '
    'deployment through GitHub Actions and AWS CodePipeline.')

# Education
set_single(p[103], '2013 to 2017, Anantapur, India')
set_single(p[104],
    'Bachelor of Technology (B.Tech.) in Computer Science and Engineering, Jawaharlal Nehru '
    'Technological University Anantapur (JNTUA)', bold=True)

# Certifications - AWS first
set_single(p[106], 'AWS Certified Solutions Architect, Associate (SAA-C03)')
set_single(p[107], 'Databricks Certified Associate Developer for Apache Spark')
set_single(p[108], 'Snowflake SnowPro Core Certification')
set_single(p[109], 'AWS Certified Data Engineer, Associate (in preparation)')

# ---- Remove unwanted lines/sections ----
# [2] contact line, [3] links line
# [111] LANGUAGES header, [112] English, [113] German
# [114] ADDITIONAL INFORMATION header, [115] Work Authorisation, [116] Availability, [117] Work Mode
# Keep [118] Other: line
for idx in (2, 3, 110, 111, 112, 113, 114, 115, 116, 117):
    el = p[idx]._element
    el.getparent().remove(el)

# ---- Global plain-ASCII scrub over every remaining run (preserves bold) ----
REPL = {
    '—': ' - ',   # em dash
    '–': ' - ',   # en dash
    '·': ', ',     # middle dot
    '→': ' to ',   # right arrow
    '←': ' ',      # left arrow
    '…': '...',    # ellipsis
    '‘': "'", '’': "'",
    '“': '"', '”': '"',
    ' ': ' ',      # nbsp
    ' ': ' ', ' ': ' ',
    '&': ' and ',
    '»': '', '«': '',
}
def scrub(t):
    for k, v in REPL.items():
        t = t.replace(k, v)
    t = t.replace('> ', 'over ').replace('>', ' over ')
    t = t.replace('< ', 'under ').replace('<', ' under ')
    t = t.replace('~', 'around ')
    t = re.sub(r' {2,}', ' ', t)
    return t

for para in doc.paragraphs:
    for run in para.runs:
        if run.text:
            run.text = scrub(run.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = scrub(run.text)

out = 'data/resumes/Jeevan-Kondasingu-AWS-Iceberg-DataEngineer-Architect.docx'
doc.save(out)
print(f'Saved: {out}')
