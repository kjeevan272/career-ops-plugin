import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('data/Jeevan_KK_Data_Engineer_DE.docx')

def set_single(para, text, bold=False, size=None):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = para.add_run(text)
    r.bold = bold
    if size: r.font.size = size

def set_two(para, t1, b1, t2, b2):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r1 = para.add_run(t1); r1.bold = b1
    r2 = para.add_run(t2); r2.bold = b2

p = doc.paragraphs

# [1] Subtitle
set_single(p[1],
    'Senior Data Engineer & Databricks Platform Engineer | Python · SQL · Spark · Power BI',
    bold=False, size=152400)

# [5] Summary line 1
set_single(p[5],
    'Senior Data Engineer with 8+ years of experience building and operating end-to-end data and analytics platforms as the technical foundation for BI, reporting, and AI solutions. Deep expertise in Databricks platform engineering — standardisation, automation, Unity Catalog governance, and continuous platform enablement across enterprise environments. Experienced connecting and processing heterogeneous operational data sources including Kafka event streams, Delta tables, relational databases, and file-based formats.')

# [6] Summary line 2
set_single(p[6],
    'Delivered high-performance data pipelines from raw ingestion through orchestrated processing to analytics-ready Delta tables — with Power BI, Databricks SQL, and API consumption layers. Proven track record in data quality frameworks, monitoring, incident analysis, and platform performance at scale. Certified: AWS Solutions Architect (Associate), Databricks Certified Developer for Apache Spark, Snowflake SnowPro Core.')

# [8] Cloud & Infrastructure — Azure/Databricks prominent
set_two(p[8], 'Cloud & Infrastructure: ', True,
    'Azure (ADF, ADLS Gen2, Azure Blob Storage, Azure DevOps, AKS, Databricks on Azure) · AWS (Glue, S3, Redshift, EMR, Kinesis, DMS, Lambda, Step Functions, Athena, CloudWatch, IAM, KMS, Lake Formation) · GCP',
    False)

# [9] Big Data — Databricks platform features + Delta + Kafka
set_two(p[9], 'Big Data & Processing: ', True,
    'Databricks (Unity Catalog, Delta Live Tables, Auto Loader, Databricks SQL, Databricks Workflows, platform standardisation & automation) · Apache Spark, PySpark · Delta Lake · Apache Kafka · Apache Flink · Apache Airflow (MWAA) · Amazon EMR · Hadoop, HDFS · Apache Iceberg',
    False)

# [10] Data Warehouses
set_two(p[10], 'Data Warehouses & Databases: ', True,
    'Delta Lake · Snowflake · Amazon Redshift · PostgreSQL · Oracle · SQL Server · MySQL · DynamoDB · Google BigQuery',
    False)

# [11] ETL — orchestration + automation angle
set_two(p[11], 'ETL & Data Integration: ', True,
    'Databricks Workflows · Delta Live Tables · Apache Airflow · dbt · AWS Glue · Great Expectations · CDC patterns · Auto Loader · AWS DMS · CI/CD pipeline automation · AppFlow',
    False)

# [15] BI & Visualisation — Power BI leading
set_two(p[15], 'BI & Visualisation: ', True,
    'Power BI (Direct Lake, Databricks SQL connector, automated report deployment) · Amazon QuickSight · Tableau · Pandas · Matplotlib · Seaborn',
    False)

# [17] Methodology — platform engineering angle
set_two(p[17], 'Methodology: ', True,
    'Agile/Scrum · Platform Engineering (standardisation, automation, enablement) · TDD · Code Reviews · Technical Mentoring · Stakeholder Management',
    False)

# [21] Role description — platform engineering + AI data foundation
set_single(p[21],
    'Designing, building, and operating end-to-end data and analytics platforms for international clients across AWS, Azure, and GCP. Acting as Databricks platform engineer: standardising pipeline patterns, automating deployments, enabling teams, and continuously developing the platform. Developing analytics-ready data products and BI reports as the technical foundation for AI and reporting solutions across multiple client engagements.')

# [42] Databricks project — reframe as platform engineering
set_two(p[42], 'Insurance / Analytics — ', True,
    'Databricks Platform Engineering — Standardisation, Automation & AI Data Foundation',
    True)

# [43] bullet 1 — platform standardisation
set_single(p[43],
    'Led Databricks platform engineering: defined reusable pipeline templates using Delta Live Tables (DLT), standardised ingestion patterns with Auto Loader, and established Unity Catalog as the governance layer for data ownership, lineage, and access control across teams.')

# [44] bullet 2 — automation + enablement
set_single(p[44],
    'Automated platform operations: CI/CD deployment pipelines for Databricks assets via Azure DevOps, templated DAG/workflow provisioning, and pre-release validation scripts — enabling teams to onboard new pipelines without manual platform intervention.')

# [45] bullet 3 — end-to-end pipelines
set_single(p[45],
    'Built end-to-end data pipelines from operational source integration (Kafka, S3, relational databases) through Medallion layers (Bronze / Silver / Gold) to analytics-ready Delta tables consumed by Power BI, Databricks SQL, and downstream AI models.')

# [46] bullet 4 — performance/monitoring
set_single(p[46],
    'Tuned PySpark workloads on Databricks (Delta cache, Z-ORDER, data skew, cluster auto-scaling) — reduced processing time by ~35% and cut infrastructure costs; set up monitoring, alerting, and incident diagnostics for stable platform operation.')

# [47] bullet 5 — reporting/BI
set_single(p[47],
    'Developed Power BI dashboards and Databricks SQL reporting modules as standardised, reusable analytics components, enabling self-service BI for operations and management stakeholders.')

# [48] bullet 6 — mentoring/architecture
set_single(p[48],
    'Led a team of 4 engineers through code reviews, architectural guidance, and platform knowledge-sharing sessions; contributed to data architecture decisions in coordination with IT architecture and specialist departments.')

# [49] Kafka project — operational data streams
set_two(p[49], 'IT Services / Streaming Analytics — ', True,
    'Real-Time Operational Data Pipeline — Kafka, Spark Structured Streaming & Event Processing',
    True)

# [50] bullet 1 — operational event ingestion
set_single(p[50],
    'Designed and implemented a Kafka-based real-time ingestion pipeline processing > 50M operational events per day at sub-second latency, integrating structured event streams alongside unstructured and semi-structured operational data formats.')

# [52] bullet 3 — monitoring/SLOs
set_single(p[52],
    'Persisted processed data into AWS S3, Redshift, and DynamoDB as Delta tables; implemented end-to-end observability (CloudWatch, Datadog) covering pipeline failures, data freshness, completeness SLOs, and incident alerting.')

# [66] Airflow project — orchestration platform
set_two(p[66], 'IT Services / Orchestration — ', True,
    'Data Pipeline Orchestration Platform — Airflow, CI/CD & Automated Operations',
    True)

# [67] bullet 1
set_single(p[67],
    'Designed and operated 40+ Airflow DAGs for automated data ingestion, transformation, and validation workflows, achieving > 99% task success rate in production — covering batch, incremental, and event-triggered pipeline patterns.')

# [68] bullet 2
set_single(p[68],
    'Implemented SLA monitoring, retry logic, alerting callbacks, and custom sensors for upstream data readiness — enabling reliable, automated orchestration of end-to-end processing jobs without manual intervention.')

# [69] bullet 3
set_single(p[69],
    'Automated DAG deployment, backfill/recovery workflows, and pre-release validation via GitHub Actions CI/CD; standardised pipeline templates to accelerate onboarding of new data sources and reporting modules.')

# [83] Great Expectations project — data quality for AI
set_two(p[83], 'Pharma / Life Sciences — ', True,
    'Data Quality & Observability Platform — Foundation for Analytics and AI',
    True)

# [84] bullet 1
set_single(p[84],
    'Defined and enforced data quality contracts across 100+ datasets using Great Expectations — validating schema consistency, completeness, referential integrity, and statistical distributions as prerequisites for AI algorithm inputs and operational reporting.')

# [85] bullet 2
set_single(p[85],
    'Built automated data profiling, freshness checks, and quality dashboards providing platform users and stakeholders with continuous visibility into data reliability and pipeline health.')

# [86] bullet 3
set_single(p[86],
    'Integrated quality checks with Airflow and Spark pipelines, blocking downstream processing on critical failures and triggering structured incident alerts with lineage context and remediation guidance.')

out = 'data/resumes/Jeevan-Kondasingu-FrankfurtAirport-SeniorDataEngineer-Databricks.docx'
doc.save(out)
print(f'Saved: {out}')
