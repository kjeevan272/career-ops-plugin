"""
Rebuild Jeevan's CV as a tight 2-page document.
- Removes LinkedIn URL
- City → "Germany"
- Projects compressed to architecture + problem + outcome bullets
- Fits within 2 A4 pages
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "data" / "Jeevan_KK_Data_Engineer_DE.docx"
OUT  = ROOT / "data" / "Jeevan_KK_Data_Engineer_DE.docx"
BAK  = ROOT / "data" / "Jeevan_KK_Data_Engineer_DE.bak2.docx"

shutil.copy(SRC, BAK)

# ── helpers ────────────────────────────────────────────────────────────────

def _set_spacing(p, before=0, after=0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    fmt.line_spacing = Pt(11)

def _add_border(paragraph, color="444444", sz="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def para(doc, before=0, after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    _set_spacing(p, before, after)
    p.alignment = align
    return p

def run(p, text, bold=False, size=9.5, italic=False, color=None):
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r

def section_header(doc, text):
    p = para(doc, before=7, after=2)
    run(p, text, bold=True, size=11)
    _add_border(p)
    return p

def competency(doc, label, value):
    p = para(doc, before=0, after=1)
    run(p, label + ": ", bold=True, size=9)
    run(p, value, size=9)
    return p

def project_title(doc, text):
    p = para(doc, before=4, after=1)
    run(p, text, bold=True, size=9.5)
    return p

def bullet(doc, text):
    p = para(doc, before=0, after=1)
    p.paragraph_format.left_indent  = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run(p, "• ", bold=False, size=9)
    run(p, text, size=9)
    return p

# ── build document ──────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4, tight margins
sec = doc.sections[0]
sec.page_width    = Inches(8.27)
sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(0.75)
sec.right_margin  = Inches(0.75)
sec.top_margin    = Inches(0.6)
sec.bottom_margin = Inches(0.6)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)

# ── HEADER ─────────────────────────────────────────────────────────────────

p = para(doc, before=0, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "JEEVAN KUMAR KONDASINGU", bold=True, size=18)

p = para(doc, before=0, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Senior Data Engineer  |  AWS · Snowflake · Databricks · PySpark", size=11)

p = para(doc, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Germany  |  +49 152 5587 9826  |  kjeevan272@gmail.com  |  github.com/kjeevan272", size=9.5)

# ── PROFESSIONAL SUMMARY ───────────────────────────────────────────────────

section_header(doc, "PROFESSIONAL SUMMARY")

p = para(doc, before=0, after=3)
run(p, (
    "Senior Data Engineer with 9+ years building cloud-native data platforms for global enterprise clients "
    "across Retail, E-Commerce, Insurance, Pharma, and Real Estate. Expert in Lakehouse architecture "
    "(Delta Lake, Apache Iceberg, Medallion pattern), real-time streaming (Kafka, Kinesis, Flink), and "
    "large-scale migrations to AWS, Snowflake, and Databricks. Delivered measurable outcomes: 40% pipeline "
    "speed gains, 60% BI query improvements, 30% cost reductions, and >99% data quality at terabyte scale "
    "across multi-cloud (AWS, Azure, GCP) environments with GDPR/HIPAA compliance."
), size=9.5)

# ── CORE COMPETENCIES ──────────────────────────────────────────────────────

section_header(doc, "CORE COMPETENCIES")

competency(doc, "Cloud & Infrastructure",
    "AWS (Glue, S3, Redshift, EMR, Kinesis, DMS, Lambda, Step Functions, Athena, CloudWatch, IAM, KMS, "
    "Lake Formation, CodePipeline) · Azure (ADF, Databricks, AKS, DevOps) · GCP")

competency(doc, "Big Data & Processing",
    "Apache Spark · PySpark · Kafka · Flink · Airflow (MWAA) · "
    "Databricks (Workflows, Auto Loader, Delta Live Tables, Unity Catalog) · EMR · Hadoop")

competency(doc, "Lakehouse & Warehouses",
    "Snowflake · Amazon Redshift · Delta Lake · Apache Iceberg · BigQuery · "
    "DynamoDB · PostgreSQL · Oracle · SQL Server")

competency(doc, "ETL & Integration",
    "AWS Glue · dbt · AWS DMS · CDC · Salesforce Bulk/REST/Streaming API · "
    "AppFlow · Great Expectations · SSIS")

competency(doc, "Engineering",
    "Python · SQL · PySpark · Spark SQL · Bash · Terraform (IaC) · "
    "Docker · Kubernetes (K8s) · Git · CI/CD (GitHub Actions, GitLab CI, CodePipeline)")

competency(doc, "Data Modelling",
    "Star/Snowflake Schema · Medallion Architecture (Bronze/Silver/Gold) · "
    "Data Products · Dimensional Modelling · Data Vault · MDM")

competency(doc, "Observability & Governance",
    "GDPR/HIPAA · RBAC · AWS IAM/KMS · Lake Formation · "
    "Data Lineage · Schema Evolution · CloudWatch · Datadog")

competency(doc, "BI & ML",
    "Power BI · Tableau · Amazon QuickSight · scikit-learn · TensorFlow · Predictive Analytics")

# ── PROFESSIONAL EXPERIENCE ────────────────────────────────────────────────

section_header(doc, "PROFESSIONAL EXPERIENCE")

# Company header
p = para(doc, before=3, after=0)
run(p, "Data Engineer  —  Global Research & Innovation Centre (UK) Limited", bold=True, size=10)

p = para(doc, before=0, after=1)
run(p, "07/2017 – Present  |  Germany  |  Multi-cloud delivery for international enterprise clients", size=9, italic=True)

# ── Project 1
project_title(doc, "Real Estate — Salesforce CRM → AWS S3 Data Lake Migration  (12 TB)")
bullet(doc,
    "Architected multi-zone Medallion Data Lake (Bronze→Silver→Gold) migrating 12 TB of Salesforce "
    "documents via CDC pipelines (Bulk/REST API + AWS DMS + Glue Workflows); 40% faster processing, "
    "30% storage cost reduction via S3 Lifecycle Policies + Glacier archival.")
bullet(doc,
    "Integrated Snowflake/Redshift external tables with Athena for self-service BI across 50+ users; "
    "AI/ML personalisation models (scikit-learn, TensorFlow) on Salesforce Experience Cloud portals; "
    "GDPR-compliant archival automation.")

# ── Project 2
project_title(doc, "Retail Analytics — Oracle & Azure Blob → Snowflake Migration  (8 TB)")
bullet(doc,
    "Migrated >8 TB Oracle/Azure Blob into Snowflake via AWS Glue (PySpark) with incremental "
    "high-watermark CDC (AWS DMS); achieved source-to-target latency <15 min across terabyte-scale "
    "fact/dimension tables.")
bullet(doc,
    "Tuned Snowflake via clustering, micro-partition pruning, and MERGE upserts — 60% BI query improvement; "
    "automated orchestration with Step Functions + Lambda + Glue Workflows; monitoring via CloudWatch "
    "and Snowflake Query History.")

# ── Project 3
project_title(doc, "E-Commerce — Enterprise Data Warehouse  (200M+ records/day)")
bullet(doc,
    "Built serverless ETL (AWS Glue + PySpark) ingesting >200M records/day from heterogeneous sources "
    "(CSV, JSON, SFTP, SharePoint via AppFlow); schema evolution, idempotent processing, and partition "
    "tuning reduced cluster costs 25% on terabyte-scale datasets.")
bullet(doc,
    "Designed Star/Snowflake schemas in Redshift with MDM for single source of truth across 6+ enterprise "
    "datasets; automated data quality framework achieving >99% quality score; semantic layer via views "
    "and stored procedures.")

# ── Project 4
project_title(doc, "Insurance — Real-Time Data Platform  (Databricks + EMR + Glue)")
bullet(doc,
    "Architected end-to-end real-time pipeline: Kinesis/S3 → Bronze Layer (Databricks Auto Loader) → "
    "PySpark transformation → Parquet/Delta Gold layer using Databricks Workflows + EMR + Glue; "
    "CDC for guaranteed source-target consistency.")
bullet(doc,
    "Tuned Spark on EMR (data skew, partitioning, resource allocation) — 35% processing time reduction; "
    "CI/CD for Databricks notebooks via GitHub Actions + CodePipeline; led team of 4 engineers with "
    "code reviews and training.")

# ── Project 5
project_title(doc, "Streaming Analytics — Kafka + Kinesis + Spark/Flink  (50M+ events/day)")
bullet(doc,
    "Designed Kafka ingestion pipeline handling >50M events/day at sub-second latency; real-time "
    "processing via Spark Structured Streaming + Apache Flink with exactly-once delivery and "
    "fault-tolerance under sustained high-throughput load.")
bullet(doc,
    "Persisted processed data to S3/Redshift/DynamoDB with full observability (CloudWatch + Datadog); "
    "engineered for high availability and data consistency under partition failures.")

# ── Project 6
project_title(doc, "Data Lakehouse — AWS S3 + Lake Formation + ADLS  (Medallion Architecture)")
bullet(doc,
    "Architected Medallion Lakehouse (Bronze/Silver/Gold) with Delta Lake + Apache Iceberg for "
    "ACID-compliant workloads; Parquet/ORC/Avro format strategy with partition pruning for "
    "maximum query throughput.")
bullet(doc,
    "Schema evolution, data lineage tracking, and cataloguing via AWS Glue Data Catalog; RBAC + "
    "encryption via Lake Formation + IAM/KMS; query optimisation with Athena, Presto, "
    "and Redshift Spectrum.")

# ── Project 7
project_title(doc, "Insurance / Clickstream — Behavioural Analytics  (Kafka + Snowflake + Tableau)")
bullet(doc,
    "End-to-end behavioural analytics pipeline: Kafka/Kinesis → Spark/Flink → Snowflake/BigQuery; "
    "dbt ELT models with Star Schema; PII compliance, data governance, and role-based access controls "
    "across web/mobile event streams.")
bullet(doc,
    "40+ Airflow DAGs for orchestration (>99% success rate); Tableau + Power BI dashboards for "
    "product and marketing teams; automated data quality and anomaly detection.")

# ── Project 8
project_title(doc, "Pharma — Data Quality & Compliance  (Great Expectations · GDPR/HIPAA)")
bullet(doc,
    "Automated data quality framework across 100+ datasets: validation rules (nulls, duplicates, "
    "schema drift), anomaly detection, and pipeline-blocking alerts integrated with Airflow + "
    "Spark + Glue; profiling reports for audit readiness.")
bullet(doc,
    "Ensured GDPR/HIPAA compliance across all pipeline stages; schema mismatch alerting reduced "
    "downstream data incidents and enabled regulatory sign-off across pharma workloads.")

# ── EDUCATION ──────────────────────────────────────────────────────────────

section_header(doc, "EDUCATION")

p = para(doc, before=2, after=0)
run(p, "B.Tech. in Computer Science & Engineering", bold=True, size=9.5)
run(p, "  —  Jawaharlal Nehru Technological University Anantapur (JNTUA), India  |  2013 – 2017", size=9.5)

# ── CERTIFICATIONS ─────────────────────────────────────────────────────────

section_header(doc, "CERTIFICATIONS")

for cert in [
    "AWS Certified Solutions Architect – Associate",
    "Databricks Certified Associate Developer for Apache Spark",
    "Snowflake SnowPro Core Certification",
    "AWS Certified Data Engineer – Associate  (in preparation)",
]:
    bullet(doc, cert)

# ── LANGUAGES ──────────────────────────────────────────────────────────────

section_header(doc, "LANGUAGES")

p = para(doc, before=2, after=0)
run(p, "English: ", bold=True, size=9.5)
run(p, "C1 Professional Working Proficiency (working language)   ", size=9.5)
run(p, "German: ", bold=True, size=9.5)
run(p, "A2 – actively developing", size=9.5)

# ── SAVE ───────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"✓ Saved: {OUT}")
print(f"  Backup: {BAK}")
