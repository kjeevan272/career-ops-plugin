import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('data/Jeevan_KK_Data_Engineer_DE.docx')

def set_single(para, text, bold=False, size=None):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = para.add_run(text)
    r.bold = bold
    if size: r.font.size = size

def set_two(para, t1, b1, t2, b2):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r1 = para.add_run(t1); r1.bold = b1
    r2 = para.add_run(t2); r2.bold = b2

def set_multi(para, parts, size=None):
    # parts = list of (text, bold) tuples — allows bolding specific words inline
    for run in para.runs:
        run._r.getparent().remove(run._r)
    for text, bold in parts:
        r = para.add_run(text)
        r.bold = bold
        if size: r.font.size = size

p = doc.paragraphs

# [1] Subtitle — Azure Databricks bold and leading
set_multi(p[1], [
    ('Senior Data Engineer | ', False),
    ('Azure Databricks', True),
    (' · PySpark · Delta Lake · Power BI · ', False),
    ('Azure DevOps', True),
], size=152400)

# [5] Summary line 1 — 70% Databricks focus
set_single(p[5],
    'Senior Data Engineer with 8+ years of experience designing, building, and operating enterprise-scale data platforms on Azure Databricks. Deep expertise in Databricks notebooks, workflows, Delta Live Tables, Unity Catalog governance, Medallion architecture (Bronze / Silver / Gold), and end-to-end pipeline delivery from raw ingestion through analytics-ready Delta tables. Specialist in PySpark performance tuning — partitioning, Z-ORDER, broadcast joins, data skew resolution, and Delta cache optimisation — delivering measurable throughput and cost improvements on large-scale datasets.')

# [6] Summary line 2 — Power BI DAX, CI/CD, ServiceNow, SQL, Azure cert
set_single(p[6],
    'Delivered production Power BI reports with DAX-driven KPI models and Direct Lake connectivity on Databricks SQL. Experienced ingesting enterprise SaaS and ITSM data (ServiceNow REST API patterns, Salesforce Bulk/REST API, CDC pipelines) into Azure Data Lake Gen2 and Databricks Lakehouse. Strong CI/CD discipline — Azure DevOps pipelines for Databricks asset deployment, automated testing, and environment promotion. Certified: Microsoft Azure Data Engineer Associate (DP-203, in progress), Databricks Certified Developer for Apache Spark, AWS Solutions Architect (Associate), Snowflake SnowPro Core.')

# [8] Cloud & Infrastructure — Azure bold and leading
set_multi(p[8], [
    ('Cloud & Infrastructure: ', True),
    ('Azure', True),
    (' (Databricks, ADLS Gen2, ADF, Azure SQL, Azure DevOps, AKS, Azure Blob Storage, Key Vault) · ', False),
    ('AWS', False),
    (' (Glue, S3, Redshift, EMR, Kinesis, Lambda, Step Functions, CloudWatch, IAM) · GCP (BigQuery, Dataflow)', False),
])

# [9] Big Data — Azure Databricks bold and dominant
set_multi(p[9], [
    ('Big Data & Processing: ', True),
    ('Azure Databricks', True),
    (' (Notebooks, Workflows, Delta Live Tables, Auto Loader, Unity Catalog, Databricks SQL, Cluster Management, Delta Cache, Z-ORDER) · Apache Spark, PySpark (transformations, joins, aggregations, window functions, performance tuning, data skew handling) · Delta Lake · Apache Kafka · Apache Flink · Apache Airflow (MWAA) · Amazon EMR · Apache Iceberg', False),
])

# [10] Data Warehouses — Azure SQL and Synapse bold
set_multi(p[10], [
    ('Data Warehouses & Databases: ', True),
    ('Azure SQL · Azure Synapse Analytics', True),
    (' · Delta Lake · Snowflake · Amazon Redshift · PostgreSQL · Oracle · SQL Server · MySQL · DynamoDB · Google BigQuery', False),
])

# [11] ETL & Integration — Azure Data Lake bold
set_multi(p[11], [
    ('ETL & Data Integration: ', True),
    ('Azure Data Lake Gen2 · ', True),
    ('Databricks Workflows · Delta Live Tables · Auto Loader · REST API ingestion (ServiceNow, Salesforce Bulk/REST/Streaming API) · CDC patterns (incremental loads, high-watermark, DMS) · dbt · AWS Glue · Great Expectations · SSIS · AppFlow', False),
])

# [12] Programming — Python + SQL explicit
set_two(p[12], 'Programming: ', True,
    'Python, PySpark, SQL (complex joins, window functions, query optimisation, execution plans), Spark SQL, DAX (Power BI measures & calculated columns), Bash/Linux, YAML',
    False)

# [14] DevOps — Azure DevOps bold and leading
set_multi(p[14], [
    ('DevOps & Governance: ', True),
    ('Azure DevOps', True),
    (' (CI/CD pipelines for Databricks notebooks, workflows, and infrastructure) · GitHub Actions · GitLab CI · Terraform (IaC) · Docker · Kubernetes · GDPR-compliant pipelines · Unity Catalog RBAC · SOX · HIPAA', False),
])

# [15] BI — Power BI DAX explicit
set_two(p[15], 'BI & Visualisation: ', True,
    'Power BI (DAX measures, calculated columns, KPI modelling, Direct Lake, Databricks SQL connector, automated report deployment) · Amazon QuickSight · Tableau · Pandas · Matplotlib · Seaborn',
    False)

# [17] Methodology — debugging + documentation explicit
set_two(p[17], 'Methodology: ', True,
    'Agile/Scrum · Debugging & Root Cause Analysis · Technical Documentation (runbooks, architecture diagrams, pipeline specs) · TDD · Code Reviews · Technical Mentoring · Stakeholder Management',
    False)

# [21] Role description — Databricks platform + enterprise ingestion
set_single(p[21],
    'Designing, building, and operating enterprise-scale data platforms on Azure Databricks and AWS for international clients. Specialised in Databricks notebook and workflow development, Delta Lake architecture, PySpark performance optimisation, and Power BI integration. Lead engineer on Medallion lakehouse builds, enterprise SaaS data ingestion projects, and CI/CD-automated pipeline deployments — with full ownership from architecture through production operations.')

# [42] Project 1 — ServiceNow/Enterprise SaaS Ingestion into Databricks (headline project)
set_two(p[42], 'Enterprise IT / ITSM — ', True,
    'ServiceNow & Enterprise SaaS Data Ingestion into Azure Databricks Lakehouse',
    True)

# [43] bullet 1 — REST API ingestion architecture
set_single(p[43],
    'Architected and implemented a scalable REST API ingestion pipeline pulling incident, CMDB, change, and service request data from enterprise ITSM and SaaS systems into Azure Data Lake Gen2 using paginated API calls, OAuth2 authentication, and incremental high-watermark CDC — processing millions of records per daily load with full audit trail.')

# [44] bullet 2 — Databricks Medallion pipeline
set_single(p[44],
    'Built end-to-end Databricks Medallion architecture (Bronze / Silver / Gold): Auto Loader for raw ingestion into Bronze Delta tables, PySpark transformation notebooks for Silver cleansing and standardisation, Gold aggregation layer serving Power BI Direct Lake reports and Databricks SQL dashboards for IT operations and management.')

# [45] bullet 3 — PySpark optimisation specifics
set_single(p[45],
    'Tuned PySpark workloads for large ITSM datasets: resolved data skew using salting on high-cardinality join keys, replaced shuffle joins with broadcast joins for dimension tables, applied Z-ORDER clustering on frequently filtered columns — reducing query execution time by ~55% and Delta scan costs significantly.')

# [46] bullet 4 — Power BI DAX
set_single(p[46],
    'Delivered Power BI dashboards with DAX-driven KPI models: calculated measures for SLA compliance rates, incident resolution time distributions, MTTR trends, and change success rates — using time intelligence DAX functions (DATEADD, CALCULATE, ALL, SAMEPERIODLASTYEAR) for period-over-period comparisons consumed by IT leadership.')

# [47] bullet 5 — CI/CD Azure DevOps
set_single(p[47],
    'Automated pipeline deployment via Azure DevOps: CI/CD pipelines for Databricks notebooks, Delta Live Tables definitions, and workflow configurations — with dev/staging/production environment promotion, automated schema validation tests, and rollback on failure; zero manual deployments to production.')

# [48] bullet 6 — documentation + debugging
set_single(p[48],
    'Authored comprehensive technical documentation: pipeline architecture diagrams, data lineage maps, debugging runbooks for common failure modes (API rate limits, schema drift, partition imbalance), and onboarding guides — enabling team members to independently diagnose and resolve production incidents.')

# [49] Project 2 — PySpark Performance Optimisation (Kafka / Streaming)
set_two(p[49], 'Insurance / Real-Time Platform — ', True,
    'PySpark Performance Engineering — Databricks, Delta Lake & Structured Streaming',
    True)

# [50] bullet 1 — performance baseline and diagnosis
set_single(p[50],
    'Diagnosed and resolved systemic PySpark performance issues on a Databricks cluster processing >200M insurance telemetry records daily: identified data skew, excessive shuffles, and unoptimised join strategies using Spark UI, DAG visualisation, and stage-level execution plan analysis.')

# [51] bullet 2 — optimisation techniques applied
set_single(p[51],
    'Applied targeted PySpark optimisations: partition pruning on Delta tables via Z-ORDER and liquid clustering, broadcast joins for lookup tables under 100MB, repartition() and coalesce() strategies to eliminate small-file problems, Delta cache warm-up for repeated scans — overall processing time reduced by ~40%, cluster cost reduced by ~35%.')

# [52] bullet 3 — Delta Lake specifics
set_single(p[52],
    'Implemented Delta Lake best practices: OPTIMIZE with Z-ORDER on query predicates, VACUUM for storage cost control, schema evolution with mergeSchema, and MERGE-based upserts for CDC — achieving ACID-compliant, idempotent pipeline runs with full time-travel capability for debugging and reprocessing.')

# [66] Project 3 — Airflow Orchestration with CI/CD
set_two(p[66], 'IT Services / Orchestration — ', True,
    'Databricks Workflow Orchestration & CI/CD Automation — Azure DevOps',
    True)

# [67] bullet 1
set_single(p[67],
    'Designed and operated 40+ Databricks Workflows and Airflow DAGs covering batch ingestion, incremental transformation, data quality validation, and Gold-layer aggregation — achieving >99% task success rate in production across multi-environment deployments.')

# [68] bullet 2
set_single(p[68],
    'Built Azure DevOps CI/CD pipelines for all Databricks assets: automated notebook linting, unit test execution, integration test runs against staging Delta tables, and environment-gated promotion (dev → staging → production) — eliminating manual deployment errors and reducing release cycle from days to hours.')

# [69] bullet 3
set_single(p[69],
    'Implemented structured debugging and incident response: Spark UI analysis for job failures, structured logging in notebooks, alerting on SLA breach and data quality gate failures, and documented runbooks for each pipeline — enabling on-call engineers to resolve incidents independently without escalation.')

# [83] Project 4 — Data Quality + SQL Optimisation (Pharma)
set_two(p[83], 'Pharma / Analytics — ', True,
    'Data Quality Framework & SQL Performance Optimisation — Great Expectations & Databricks SQL',
    True)

# [84] bullet 1 — data quality gates
set_single(p[84],
    'Implemented Great Expectations data quality contracts across 100+ datasets as blocking gates in Databricks Workflows — validating schema consistency, null rates, referential integrity, and statistical distributions before downstream Gold-layer processing; zero data quality incidents reaching production reporting.')

# [85] bullet 2 — SQL optimisation specifics
set_single(p[85],
    'Optimised complex Databricks SQL and PySpark SQL queries: rewrote correlated subqueries as window functions, replaced CROSS JOINs with broadcast joins, added partition filters to eliminate full-table scans, and applied predicate pushdown — reducing average dashboard query time by ~60% and Databricks SQL warehouse credit consumption significantly.')

# [86] bullet 3 — documentation and compliance
set_single(p[86],
    'Documented data models, SQL query libraries, and validation rule catalogues for regulatory compliance (GDPR, HIPAA) — enabling audit evidence generation in under one hour and providing self-service SQL reference for analyst and data science teams.')

# Certifications — Azure first, then Databricks
set_single(p[106], 'Microsoft Certified: Azure Data Engineer Associate (DP-203) — in preparation')
set_single(p[107], 'Databricks Certified Associate Developer for Apache Spark')
set_single(p[108], 'AWS Certified Solutions Architect – Associate')
set_single(p[109], 'Snowflake SnowPro Core Certification')

out = 'data/resumes/Jeevan-Kondasingu-AzureDatabricks-SeniorDataEngineer.docx'
doc.save(out)
print(f'Saved: {out}')
