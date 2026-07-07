import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document('data/Jeevan_KK_Data_Engineer_DE.docx')

def set_single(para, text, bold=False, size=None):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r = para.add_run(text)
    r.bold = bold
    if size: r.font.size = size

def set_two(para, t1, b1, t2, b2):
    for run in para.runs:
        run._r.getparent().remove(run._r)
    r1 = para.add_run(t1); r1.bold = b1
    r2 = para.add_run(t2); r2.bold = b2

p = doc.paragraphs

# [1] Subtitle — multi-cloud consulting focus
set_single(p[1],
    'Senior Cloud Data Engineer | AWS · Azure · GCP · Databricks · Snowflake · PySpark',
    bold=False, size=152400)

# [5] Summary — consulting + multi-cloud + end-to-end delivery
set_single(p[5],
    'Senior Cloud Data Engineer with 8+ years of experience delivering end-to-end data platform projects for clients across Retail, E-Commerce, Insurance, Real Estate, and Pharma. Proven expertise building Data Lakes, Lakehouses, and Data Warehouses on AWS, Azure, and GCP, combining Databricks, Snowflake, Apache Spark/Flink, and Delta/Iceberg into scalable, production-ready architectures. Experienced in client-facing engagements from requirements and architecture to productive use, with a strong consulting delivery track record.')

# [6] Summary line 2 — certs + IaC + architecture
set_single(p[6],
    'Holds AWS Certified Solutions Architect (Associate), Databricks Certified Developer for Apache Spark, and Snowflake SnowPro Core certifications — directly aligned with msg\'s platform focus. Delivered measurable outcomes: 40% reduction in pipeline processing time, 30% cloud cost savings, and 99%+ data quality scores. Experienced with CI/CD, Terraform IaC, and MLOps data pipeline integration.')

# [8] Cloud & Infrastructure — all three clouds prominent
set_two(p[8], 'Cloud & Infrastructure: ', True,
    'AWS (Glue, S3, Redshift, EMR, Kinesis, DMS, Lambda, Step Functions, Athena, CloudWatch, IAM, KMS, Lake Formation, CodePipeline) · Azure (ADF, ADLS Gen2, Databricks on Azure, AKS, DevOps) · GCP (BigQuery, Dataflow, Pub/Sub)',
    False)

# [9] Big Data — Databricks, Spark/Flink, Delta/Iceberg
set_two(p[9], 'Big Data & Processing: ', True,
    'Databricks (Delta Live Tables, Unity Catalog, Databricks SQL) · Apache Spark, PySpark · Apache Flink · Delta Lake · Apache Iceberg · Apache Kafka · Apache Airflow (MWAA) · Amazon EMR · Hadoop, HDFS',
    False)

# [10] Data Warehouses — Snowflake + Databricks prominent
set_two(p[10], 'Data Warehouses & Databases: ', True,
    'Snowflake · Databricks Lakehouse · Amazon Redshift · Delta Lake · Apache Iceberg · PostgreSQL · Oracle · SQL Server · MySQL · DynamoDB · Google BigQuery',
    False)

# [11] ETL — breadth of platform tooling
set_two(p[11], 'ETL & Data Integration: ', True,
    'AWS Glue · dbt · Databricks Workflows · Great Expectations · CDC patterns · AWS DMS · SSIS · AppFlow · Salesforce Bulk/REST/Streaming API',
    False)

# [14] DevOps — Terraform, CI/CD, MLOps
set_two(p[14], 'DevOps & Governance: ', True,
    'Terraform (IaC) · Git · GitHub Actions · GitLab CI · CI/CD · Docker · Kubernetes · MLOps (pipeline integration for scikit-learn, TensorFlow) · GDPR-compliant pipelines · SOX · RBAC',
    False)

# [16] ML/AI — frame as MLOps/data pipelines
set_two(p[16], 'ML/AI & MLOps: ', True,
    'scikit-learn · TensorFlow (data pipeline integration) · MLOps patterns (feature engineering pipelines, model data prep, experiment tracking) · Predictive Analytics · KNN, SVM, Random Forest',
    False)

# [21] Role description — consulting framing
set_single(p[21],
    'Delivering end-to-end cloud data platform projects for international clients across AWS, Azure, and GCP. Responsibilities span requirements analysis, architecture design, pipeline implementation, and productive go-live — advising clients from first concept through to operational platform. Leading junior engineers, contributing to team knowledge exchange, and positioning data architecture best practices.')

# Certifications — reorder to put the most relevant first and add context
# [106] AWS SA
set_single(p[106], 'AWS Certified Solutions Architect – Associate (SAA-C03)')
# [107] Databricks
set_single(p[107], 'Databricks Certified Associate Developer for Apache Spark 3.0')
# [108] Snowflake
set_single(p[108], 'Snowflake SnowPro Core Certification')
# [109] AWS Data Engineer
set_single(p[109], 'AWS Certified Data Engineer – Associate (in preparation)')

out = 'data/resumes/Jeevan-Kondasingu-msg-CloudDataEngineer.docx'
doc.save(out)
print(f'Saved: {out}')
